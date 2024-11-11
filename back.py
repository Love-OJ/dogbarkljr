from flask import Flask, request, jsonify
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from flask_cors import CORS

app = Flask(__name__)
CORS(app) 
CORS(app, origins=['http://localhost:5173'])

# 指定保存Word文档的文件夹
SAVE_FOLDER = 'saved_resumes'  # 你可以根据需要修改这个路径

# 确保保存文件夹存在
if not os.path.exists(SAVE_FOLDER):
    os.makedirs(SAVE_FOLDER)

def add_styled_paragraph(doc, text, font_size=12, bold=False, color=None, alignment=None):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment:
        paragraph.alignment = alignment
    return paragraph

def process_form_data(form):
    """直接返回表单数据字典"""
    return {
        'name': form.get('name', ''),
        'student_id': form.get('id', ''),
        'class_name': form.get('class', ''),
        'school': form.get('school', ''),
        'english_score': form.get('englishScore', ''),
        'gender': form.get('gender', ''),
        'dob': form.get('dob', ''),
        'phone': form.get('phone', ''),
        'email': form.get('email', ''),
        'introduction': form.get('introduction', ''),
        'join_student_union': form.get('joinStudentUnion', ''),
        'student_union_name': form.get('studentUnionName', ''),
        'main_tasks': form.get('mainTasks', ''),
    }


@app.route('/submit', methods=['POST'])
def submit():
    try:
        # 处理表单数据
        #print(request.json)
        data = process_form_data(request.json)
        
        # 打印调试信息
        print(data)  # 打印所有表单数据以供调试

        # 创建Word文档
        doc = Document()
        doc.add_heading('简历', level=1).bold = True

        # 添加个人信息
        add_styled_paragraph(doc, f'姓名: {data["name"]}', font_size=14, bold=True, color=(0, 0, 255))
        add_styled_paragraph(doc, f'学号: {data["student_id"]}')
        add_styled_paragraph(doc, f'班级: {data["class_name"]}')
        add_styled_paragraph(doc, f'高中母校: {data["school"]}')
        add_styled_paragraph(doc, f'高考英语成绩: {data["english_score"]}')
        add_styled_paragraph(doc, f'性别: {data["gender"]}')
        add_styled_paragraph(doc, f'出生日期: {data["dob"]}')
        add_styled_paragraph(doc, f'联系电话: {data["phone"]}')
        add_styled_paragraph(doc, f'电子邮箱: {data["email"]}')
        add_styled_paragraph(doc, f'自我介绍: {data["introduction"]}', alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)

        # 学生会信息
        add_styled_paragraph(doc, f'是否加入学生会: {data["join_student_union"]}')
        
        if data["join_student_union"].lower() == 'yes':
            add_styled_paragraph(doc, f'加入的学生会: {data["student_union_name"]}', font_size=12, bold=True)
            add_styled_paragraph(doc, f'主要任务: {data["main_tasks"]}', alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)

        # 保存Word文档
        file_name = f"{data['name']}_{data['student_id']}.docx"  # 生成文件名
        file_path = os.path.join(SAVE_FOLDER, file_name)
        doc.save(file_path)
        return jsonify(message=f"提交成功! 文档已保存到: {file_path}")
    except Exception as e:
        return jsonify(message=f"提交失败: {str(e)}"), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)
